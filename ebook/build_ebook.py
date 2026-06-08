"""
Build the publishable MS Word manuscript for:
  "Build a World-Class Website with AI"
Run:  python ebook/build_ebook.py
Out:  ebook/Build-a-World-Class-Website-with-AI.docx

Design goals
------------
* Amazon KDP / print-ready: 6x9 trim, mirror margins, front matter in roman
  numerals, body in arabic, auto Table of Contents field, page numbers.
* Reader-friendly for NON-IT readers: plain language, callout boxes
  (User Instruction / Note / Tip / Why it matters), embedded brand diagrams.
* Every build step counted sequentially (final results only; reworks omitted).
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
FIG = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "Build-a-World-Class-Website-with-AI.docx")

# Brand palette
BLUE = RGBColor(0x00, 0x73, 0xCF)
LIME = RGBColor(0x2E, 0x8B, 0x2E)   # darkened lime for readable text
POPPY = RGBColor(0xB8, 0x86, 0x00)  # darkened poppy for readable text
INK = RGBColor(0x0B, 0x1B, 0x2B)
SLATE = RGBColor(0x44, 0x55, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Callout themes: (accent text color, fill hex)
THEMES = {
    "user":  (BLUE,  "EAF3FB", "USER INSTRUCTION"),
    "note":  (SLATE, "F2F5F8", "PLAIN-ENGLISH NOTE"),
    "tip":   (LIME,  "EFF8EE", "TIP"),
    "why":   (POPPY, "FFF8E6", "WHY IT MATTERS"),
}

YEAR = "2026"

# ── low-level helpers ────────────────────────────────────────────────────────


def shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), fill)
    el.append(sh)


def cell_border(cell, color="BBBBBB", sz="8", left_accent=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        if edge == "left" and left_accent:
            e.set(qn("w:sz"), "24")
            e.set(qn("w:color"), left_accent)
        else:
            e.set(qn("w:sz"), sz)
            e.set(qn("w:color"), color)
        e.set(qn("w:space"), "0")
        borders.append(e)
    tcPr.append(borders)


def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{name}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def page_number_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def set_section_pgnum(section, fmt="decimal", start=None):
    sectPr = section._sectPr
    old = sectPr.find(qn("w:pgNumType"))
    if old is not None:
        sectPr.remove(old)
    p = OxmlElement("w:pgNumType")
    p.set(qn("w:fmt"), fmt)
    if start is not None:
        p.set(qn("w:start"), str(start))
    sectPr.append(p)


def add_toc(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click here and choose “Update Field” to build the Table of Contents."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2); run._r.append(t); run._r.append(f3)


# ── document setup ───────────────────────────────────────────────────────────

doc = Document()

# base font
normal = doc.styles["Normal"]
normal.font.name = "Georgia"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25

def style_heading(name, size, color, before, after, bold=True):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

style_heading("Title", 30, INK, 0, 6)
style_heading("Heading 1", 22, BLUE, 18, 8)
style_heading("Heading 2", 16, INK, 14, 6)
style_heading("Heading 3", 12.5, SLATE, 10, 4)

# 6x9 trim with mirrored margins
sec = doc.sections[0]
sec.page_width = Inches(6)
sec.page_height = Inches(9)
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.6)
sec.gutter = Inches(0.15)

# ── content helpers ──────────────────────────────────────────────────────────

STEP = {"n": 0}


def p(text="", align=None, size=None, color=None, italic=False, bold=False, space_after=None):
    par = doc.add_paragraph()
    if align:
        par.alignment = align
    run = par.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if space_after is not None:
        par.paragraph_format.space_after = Pt(space_after)
    return par


def h1(text):
    doc.add_paragraph(text, style="Heading 1")


def h2(text):
    doc.add_paragraph(text, style="Heading 2")


def h3(text):
    doc.add_paragraph(text, style="Heading 3")


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def pagebreak():
    doc.add_page_break()


def step(title, body):
    STEP["n"] += 1
    par = doc.add_paragraph(style="Heading 2")
    r1 = par.add_run(f"Step {STEP['n']}.  ")
    r1.font.color.rgb = LIME
    r2 = par.add_run(title)
    r2.font.color.rgb = INK
    p(body)


def callout(kind, text, title=None):
    color, fill, label = THEMES[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    cell_border(cell, color="DDDDDD", left_accent="%02X%02X%02X" % (color[0], color[1], color[2]))
    set_cell_margins(cell)
    cell.paragraphs[0].text = ""
    head = cell.paragraphs[0]
    hr = head.add_run((title or label).upper())
    hr.bold = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = color
    head.paragraph_format.space_after = Pt(3)
    body = cell.add_paragraph()
    br = body.add_run(text)
    br.font.size = Pt(10.5)
    br.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def config(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.2)
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = INK
    shade(par._p.get_or_add_pPr(), "F4F6F8")


def figure(name, caption):
    path = os.path.join(FIG, name)
    if not os.path.exists(path):
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Inches(4.7))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = SLATE


def screenshot_placeholder(fignum, what, where):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), "FAFBFC")
    cell_border(cell, color="C9D2DC")
    set_cell_margins(cell, top=260, bottom=260)
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(f"[ SCREENSHOT {fignum} ]\nCapture: {what}\nWhere: {where}")
    r.font.size = Pt(10)
    r.font.color.rgb = SLATE
    r.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════════════════════════════════
# FRONT MATTER
# ════════════════════════════════════════════════════════════════════════════

# Half-title
p()
p("BUILD A WORLD-CLASS WEBSITE WITH AI", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, color=INK, bold=True)
pagebreak()

# Cover image
fc = os.path.join(FIG, "fig_cover.png")
if os.path.exists(fc):
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run().add_picture(fc, width=Inches(4.2))
pagebreak()

# Title page
for _ in range(2):
    p()
p("Build a World-Class Website with AI", align=WD_ALIGN_PARAGRAPH.CENTER, size=26, color=INK, bold=True)
p("The Non-Coder’s Guide — From Idea to AI Chatbot, Step by Step",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=BLUE, italic=True)
p()
p("How a Bangladeshi engineering firm built a premium, AI-ready website — "
  "without writing a single line of code.", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=SLATE, italic=True)
for _ in range(4):
    p()
p("Md. Sohel Sikder", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, color=INK, bold=True)
p("Published by Sustech Technology Ltd", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=SLATE)
pagebreak()

# Copyright page
p(f"Copyright © {YEAR} Md. Sohel Sikder and Sustech Technology Ltd.", size=10)
p("All rights reserved.", size=10)
p("No part of this publication may be reproduced, stored in a retrieval system, or "
  "transmitted in any form or by any means — electronic, mechanical, photocopying, "
  "recording, or otherwise — without the prior written permission of the publisher, "
  "except for brief quotations in a review.", size=10)
p()
p("Published by Sustech Technology Ltd, Dhaka, Bangladesh.", size=10)
p(f"First edition, {YEAR}.", size=10)
p()
p("ISBN (ebook / Kindle):  uses an Amazon ASIN — see Appendix B.", size=10)
p("ISBN (paperback):  979-8-_____-_____-_   [assign before publishing]", size=10)
p("ISBN (hardcover):  979-8-_____-_____-_   [assign before publishing]", size=10)
p()
p("Disclaimer: This book documents one real project. Tool names, product names, and "
  "trademarks belong to their respective owners and are used for identification only. "
  "The information is provided “as is,” without warranty of any kind. Neither the author "
  "nor the publisher is liable for any loss arising from the use of this book.", size=9.5, color=SLATE)
p()
p("Cover and interior diagrams designed in the Sustech brand palette: "
  "True Blue, Lime Green, and Golden Poppy.", size=9.5, color=SLATE)
pagebreak()

# Dedication
for _ in range(4):
    p()
p("For every founder, marketer, and curious professional who was told "
  "“you need to be a developer for this.”\nYou don’t.",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=INK, italic=True)
pagebreak()

# Epigraph
for _ in range(3):
    p()
p("“Design for two readers at once: the human who judges in fifty milliseconds, "
  "and the machine that decides whether to cite you.”",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=12.5, color=BLUE, italic=True)
p("— The governing principle of this project",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=SLATE)
pagebreak()

# Table of contents
h1("Contents")
add_toc(doc.add_paragraph())
pagebreak()

# How to read this book
h1("How to Read This Book")
p("This book tells the true story of building a complete, modern company website — "
  "from an empty folder to a live site with an AI chatbot that can read photos. It is "
  "written for people who are NOT software developers: business owners, marketers, "
  "operations managers, and anyone learning to direct AI tools to do real work.")
p("You will not be asked to memorise code. Where technical things appear, they are "
  "explained in plain language and boxed off, so you can read straight through the story "
  "and still understand every decision.")
h3("The four kinds of boxes you’ll see")
callout("user", "The actual, human instruction that started a step — the plain request a "
        "non-technical person gave, in everyday words. This is the heart of “AI mastering”: "
        "you lead with intent, the AI handles the mechanics.")
callout("note", "A short, plain-English explanation of a technical term or idea, so you’re "
        "never lost — even if you’ve never seen the word before.")
callout("tip", "A practical shortcut, habit, or lesson learned that you can reuse on your "
        "own project.")
callout("why", "Why a step mattered — the business or quality reason behind a technical "
        "decision, so the “so what?” is always clear.")
p("Each construction step is numbered (Step 1, Step 2, …) so you can see the project grow "
  "in a clean, logical order. Reworks and dead-ends have been left out on purpose — what "
  "you see is the path that actually worked.")
pagebreak()

# Who this book is for
h1("Who This Book Is For")
bullet("Founders and business owners who want a premium website but don’t want to learn to code.")
bullet("Marketers and content managers who will run the site day to day.")
bullet("Professionals teaching themselves to direct AI tools — “AI mastering” without programming.")
bullet("Agencies and consultants who want a proven, repeatable blueprint for a modern build.")
bullet("Students and the simply curious who want to see how a real, modern website is made.")
p()
callout("note", "“AI mastering” in this book means learning to give clear instructions, make "
        "good decisions, and review results — while an AI assistant does the hands-on building. "
        "You stay the director; the AI is the crew.")
pagebreak()

# ════════════════════════════════════════════════════════════════════════════
# BODY — switch to arabic page numbers in a new section
# ════════════════════════════════════════════════════════════════════════════

# footer for the (front matter) first section: roman numerals
fm_footer = doc.sections[0].footer
fm_footer.is_linked_to_previous = False
fp = fm_footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
page_number_field(fp)
set_section_pgnum(doc.sections[0], "lowerRoman", start=1)

# New section for the body
body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
body_sec.page_width = Inches(6); body_sec.page_height = Inches(9)
body_sec.top_margin = Inches(0.75); body_sec.bottom_margin = Inches(0.75)
body_sec.left_margin = Inches(0.75); body_sec.right_margin = Inches(0.6); body_sec.gutter = Inches(0.15)
body_sec.footer.is_linked_to_previous = False
bfp = body_sec.footer.paragraphs[0]
bfp.alignment = WD_ALIGN_PARAGRAPH.CENTER
page_number_field(bfp)
set_section_pgnum(body_sec, "decimal", start=1)

# ── PART I ───────────────────────────────────────────────────────────────────
h1("Part I — The Idea and the Ground Rules")
p("Before a single page was built, the project had a clear vision and a strict set of "
  "rules. This part explains both — in plain language — because every later decision flows "
  "from them.")

h2("Chapter 1 · The Vision: A Website for Two Readers")
p("Sustech Technology Ltd is an engineering firm in Bangladesh. It designs and builds solar "
  "power systems, electrical infrastructure, lightning protection, and smart building "
  "systems for factories and large commercial clients. The old website did not reflect the "
  "quality of the work. The goal was a new site that would feel premium and trustworthy to "
  "industrial buyers — and, just as importantly, be perfectly readable by AI engines.")
figure("fig_two_readers.png", "Figure 1.1 — Every page must satisfy a human reader and a machine reader at the same time.")
p("Here is the idea that shaped everything: every page is read twice. A human visitor judges "
  "the site in about fifty milliseconds — they want beauty, clarity, and trust. A machine — "
  "Google, or an AI assistant like ChatGPT or Claude — reads the same page to decide whether "
  "to recommend the company. These two readers never conflict if you follow one discipline: "
  "the real content lives in clean, server-delivered text, and the beauty is layered on top.")
callout("why", "When someone asks an AI assistant “who can install solar at my factory in "
        "Bangladesh?”, the AI can only recommend companies whose websites it can actually read. "
        "A beautiful site the machine can’t read is invisible at the exact moment a buyer is "
        "choosing. This is the new frontier of being found online — sometimes called GEO or AEO "
        "(optimising for AI engines and answers).")
callout("note", "“Server-delivered text” means the words are baked into the page before it "
        "reaches the browser — not drawn later by hidden scripts. That is what makes a page "
        "fast for humans and fully readable for machines.")

h2("Chapter 2 · The Ground Rules and the Toolkit")
p("A written brief set the rules for the whole build. Two rules mattered above all others.")
h3("Rule 1 — Content lives in clean, readable text")
p("No important words are ever hidden inside images or scripts. This keeps the site fast, "
  "accessible, and citable by AI.")
h3("Rule 2 — The site is fully run from a control panel (CMS)")
p("The code provides the engine, the design, and the building-block types. All the actual "
  "content — menus, pages, projects, team members, photos, contact details — is managed by a "
  "non-technical admin from a private control panel. Adding a page or a menu item never "
  "requires a programmer.")
figure("fig_cms_model.png", "Figure 2.1 — Code is the engine; the CMS control panel owns the content.")
callout("note", "“CMS” stands for Content Management System — a private admin website where "
        "you log in and edit your public site’s content with forms and buttons, no code. This "
        "project uses one called Payload. Think of it as the steering wheel for your website.")
h3("The toolkit, in plain words")
bullet("Next.js — the modern framework that builds the actual website pages.")
bullet("Payload CMS — the control panel where content is managed.")
bullet("PostgreSQL — the database that stores the content safely.")
bullet("n8n — a visual automation tool that powers the AI chatbot’s brain.")
bullet("A VPS — a rented online computer where the finished site lives.")
callout("tip", "You do not need to understand these tools to follow the book. You only need "
        "to know what each one is for — exactly as a project director knows what each member of "
        "the crew does, without doing their job for them.")

h2("Chapter 3 · The Quality Discipline")
p("Professional results come from a habit, not a one-time effort. Every change followed the "
  "same loop before it was accepted as “done.”")
figure("fig_quality_loop.png", "Figure 3.1 — Nothing was called “done” until every automatic check passed.")
p("Write the change, then let the computer check it five ways: does the code make sense, is it "
  "tidy, is it formatted consistently, do the automated tests pass, and does the whole site "
  "build successfully? If any check failed, it was fixed and the loop ran again. Only a fully "
  "green result counted.")
callout("why", "This discipline is why the final site is fast and reliable. Skipping checks is "
        "how websites quietly accumulate bugs that surface in front of customers later.")
callout("tip", "Even as a non-coder directing an AI, you can insist on this: “don’t tell me "
        "it’s done until all the checks pass.” It is the single most valuable instruction you "
        "can give.")
pagebreak()

# ── PART II ──────────────────────────────────────────────────────────────────
h1("Part II — Laying the Foundation")
p("With the vision and rules set, construction began. The first three steps created the "
  "skeleton everything else would hang on.")

step("Set up the project and the quality gates",
     "An empty project was created and immediately wired with the automatic checks from "
     "Chapter 3. This sounds boring, but doing it first means quality is built in from minute "
     "one rather than bolted on at the end.")
callout("tip", "Set your standards before you build, not after. It is far cheaper to keep a "
        "house clean than to clean it once it’s a mess.")

step("Build the design system and the home-page shell",
     "Before any real content, a “design system” was created: the official fonts, colours, "
     "spacing, and reusable pieces like buttons and cards. Then the home page’s basic frame "
     "(header, main area, footer) was put in place.")
callout("note", "A “design system” is a small rulebook of fonts, colours, and components so "
        "every page looks like it belongs to the same brand — instead of each page being styled "
        "by guesswork.")

step("Turn the site into a fully content-managed CMS",
     "The Payload control panel was added and connected to the database. From this point on, "
     "menus, pages, and content blocks could be created and arranged by a non-technical admin "
     "— exactly as Rule 2 required.")
callout("why", "This is the step that makes the site sustainable. The marketing team can run it "
        "for years without calling a developer for everyday changes.")
pagebreak()

# ── PART III ─────────────────────────────────────────────────────────────────
h1("Part III — Building the Pages")
p("With the foundation ready, the site filled out page by page. Each step below added a real, "
  "visible part of the website.")

step("Add the fifth service line: Testing, Inspection & Consultancy",
     "The company’s services were modelled as content, and a fifth offering was added alongside "
     "solar, electrical, lightning protection, and smart systems.")
step("Build the project importer (spreadsheet → draft projects)",
     "Rather than typing in years of past projects by hand, a small tool was built to read an "
     "Excel spreadsheet and create draft project entries in the control panel automatically.")
callout("tip", "When you have lots of existing data, importing beats retyping — fewer errors and "
        "hours saved. Ask your AI assistant whether an import is possible before you start typing.")
step("Enrich the draft projects with sectors and factual summaries",
     "Each imported project was tagged by industry sector and given a short, accurate summary, "
     "so visitors and search engines could understand them at a glance.")
step("Create the Projects index with a sector filter",
     "A page was built that lists all projects, with a filter so a visitor can show only the "
     "sectors relevant to them (for example, only garment factories).")
step("Add the Team collection and the About page",
     "A place to manage team members was created, and a starter About page was assembled from "
     "content blocks.")
step("Add individual Service detail pages",
     "Each service got its own dedicated page with room for descriptions, images, and proof.")
step("Build the Request-a-Quote / Contact flow",
     "Because this is a business that wins work through enquiries — not online shopping — the "
     "main action across the site is “Request a Quote.” A guided enquiry form was built.")
callout("why", "Knowing your “conversion goal” shapes the whole site. For Sustech it is a "
        "quote request, never a shopping cart. Every page gently leads there.")
step("Add Solutions (industry sector) detail pages",
     "Pages were created for each industry the company serves, so a factory owner can find a "
     "page that speaks directly to their world.")
step("Add the Knowledge hub and the Capabilities page",
     "A library of helpful articles and a capabilities overview were added — useful for "
     "visitors and excellent for being found by search engines and AI.")
step("Add a dedicated Contact page",
     "Separate from the quote form, a simple contact page gave visitors phone, email, and "
     "location at a glance.")
step("Lay the SEO foundation (sitemap, robots, llms.txt, schema)",
     "Behind-the-scenes files were added that tell search engines and AI assistants how to read "
     "the site, which pages exist, and what the company is.")
callout("note", "These files are invisible to visitors but vital to machines. “Schema” is a "
        "structured summary of each page that AI and Google read like a fact sheet. “llms.txt” "
        "is a newer file that points AI assistants to your most important pages.")
step("Fill in the real company data",
     "Placeholder details were replaced with Sustech’s real address, phone numbers, and company "
     "facts — no invented statistics, ever.")
callout("tip", "Never publish fake numbers (“500+ happy clients!”) you can’t back up. AI engines "
        "and savvy buyers both punish exaggeration. Real, modest, and true beats impressive and false.")
step("Publish the About page with real content and seed the team",
     "The About page was finished with genuine company history and real team members.")
step("Add multi-dimension filters to the projects page",
     "Visitors could now filter projects by sector, service, and year together — finding exactly "
     "the proof they care about.")
step("Adopt the 10-service / 10-sector model",
     "The site’s menu of services and industries was organised into a clean ten-and-ten "
     "structure, matching how the company actually talks about its work.")
step("Seed eight real, named showcase projects",
     "Eight flagship projects were published with real names and details to build credibility.")
step("Seed five knowledge articles",
     "Five helpful, in-depth articles were published to start the knowledge library and attract "
     "search and AI traffic.")
step("Add the footer trust strip and partner bar",
     "The bottom of every page gained credibility markers and partner logos — quiet, constant "
     "reassurance.")
step("Add the WhatsApp button and a home-page knowledge preview",
     "A floating WhatsApp button made contact effortless, and the home page began previewing "
     "the latest articles.")
step("Add the home hero image and the social-share image",
     "A striking main image was added to the top of the home page, plus the picture that appears "
     "when the site is shared on social media.")
step("Add the animated statistics strip",
     "A tasteful, animated band showed the company’s real numbers — projects delivered, clients "
     "served, sectors covered, years in business.")
pagebreak()

# ── PART IV ──────────────────────────────────────────────────────────────────
h1("Part IV — Making It Premium")
p("A functional site is not the same as a world-class one. This part is where the project "
  "crossed from “correct” to “impressive.”")

step("Add the motion and animation suite",
     "Subtle, tasteful animations were introduced — content that gently fades in as you scroll. "
     "Crucially, the motion was built so it never slows the site or harms readability for "
     "machines, and it switches off for visitors who prefer reduced motion.")
callout("why", "Good motion feels premium and guides the eye. Bad motion feels cheap and "
        "distracting. The difference is restraint and respecting the visitor’s preferences.")
step("Wire the brand logo throughout the site",
     "The official Sustech logo was placed consistently across the header, footer, icons, and "
     "the machine-readable fact sheets.")
step("Add AI-generated media (service heroes, products, hero video, team photos)",
     "Professional images and a hero video were produced with AI tools and placed across the "
     "service pages and home page — premium visuals without an expensive photo shoot.")
callout("tip", "AI image and video tools can produce striking, on-brand visuals at a fraction of "
        "traditional cost — a real superpower for a non-designer directing the work.")
step("Add the BESS and LPS explainer videos",
     "Two short explainer videos — for battery storage and lightning protection — were added to "
     "their service pages, loading only when a visitor chooses to play them so the page stays fast.")
step("Unify the hover-reveal cards",
     "Project, product, and team cards were given a consistent, elegant behaviour: a clean "
     "summary that reveals more detail when you hover, focus, or tap.")
step("Adopt the True Blue / Lime Green / Golden Poppy palette",
     "The three official brand colours were applied site-wide through the design system, giving "
     "everything a coherent, confident identity.")
callout("note", "Sustech’s colours: True Blue (#0073CF), Lime Green (#32CD32), Golden Poppy "
        "(#FCC200). Defining colours once, centrally, means a single change updates the whole site.")
step("Add the dynamic per-block style system",
     "This was a standout feature: every content block gained a style panel in the control "
     "panel. A non-technical admin can now change a section’s colour scheme, width, spacing, "
     "border, and animation — all without touching code.")
callout("why", "This is Rule 2 taken to its peak: not only the content but the look of each "
        "section is now in the admin’s hands. The site can be restyled for years without a developer.")
pagebreak()

# ── PART V ───────────────────────────────────────────────────────────────────
h1("Part V — Hardening and Going Live")
p("A great site that is slow, insecure, or invisible to search has failed. This part made the "
  "site safe, fast, findable — and then put it on the internet.")

step("Run a full security, performance, and architecture audit",
     "The whole project was reviewed for security holes, speed problems, and structural "
     "weaknesses, and the important issues were fixed. The result: top-tier security headers, "
     "excellent speed scores, and a clean architecture.")
callout("note", "“Security headers” are quiet instructions the site sends to every browser that "
        "make common attacks much harder. Visitors never see them, but they protect everyone.")
step("Add the first web chat widget connected to n8n",
     "An initial chat widget was added that captured leads and connected to the n8n automation "
     "tool — the seed that would later grow into the full AI assistant.")
step("Strengthen AI-findability (GEO): llms.txt route and richer schema",
     "The machine-readable layer was upgraded so AI assistants could understand and cite the "
     "site even more reliably.")
step("Add the Hermes AI content agent pipeline",
     "An automated content assistant, “Hermes,” was set up to draft news and articles on a "
     "schedule. Importantly, Hermes can only create drafts — a human still approves and publishes.")
callout("why", "Automation should accelerate people, not replace their judgement. Letting an AI "
        "draft but never auto-publish keeps a human firmly in control of what the world sees.")
step("Make the beta site private from search (noindex)",
     "While the site lived on a test address, it was deliberately hidden from Google so the "
     "unfinished version could never appear in search results. A single switch flips it public at launch.")
step("Fix the security policy for the live build, plus accessibility and housekeeping",
     "A subtle but critical fix corrected the site’s content-security rules so the live, "
     "high-speed version worked perfectly, and accessibility was tightened.")
step("Make unknown web addresses return a proper “not found”",
     "Mistyped or dead links now return a correct “page not found” instead of a misleading blank "
     "success — better for visitors and for search engines.")
step("Set up the super-admin and a batch of refinements",
     "A secure top-level administrator account was created, along with a batch of polish: an "
     "image carousel, more style options, and refined icons.")
step("Add smooth page-to-page transitions",
     "Moving between pages was given a gentle animated fade, so navigation feels instant and "
     "premium even in the split second a new page loads.")
step("Split the navigation menus (label opens the index; arrow opens the dropdown)",
     "The Services and Solutions menus were made smarter: clicking the word opens the full list "
     "page, while clicking the little arrow opens the quick dropdown.")
step("Group the team section by category automatically",
     "The team area learned to organise people into clear groups — Leadership, Engineering, "
     "Consultants, Advisors — based on a simple choice in the control panel.")
step("Deploy the site to a live server at beta.sustechltd.com",
     "The finished site was published to a rented online computer (a VPS), kept running "
     "automatically, served securely over HTTPS, and placed behind a professional web server — "
     "while the main domain stayed safely on its existing host.")
figure("fig_architecture.png", "Figure 5.1 — The full system, now live on its own server.")
callout("note", "“VPS” = Virtual Private Server, a rented computer in a data centre that runs "
        "your site around the clock. “HTTPS” is the padlock in the address bar — an encrypted, "
        "trusted connection.")
pagebreak()

# ── PART VI ──────────────────────────────────────────────────────────────────
h1("Part VI — The Finale: An AI Chatbot That Sees")
p("The final chapter of the build is the feature most people notice first: a smart assistant in "
  "the corner of the screen. This is where the project went from an excellent website to a "
  "genuinely modern one.")

h2("Chapter · From Fixed Buttons to a Real Conversation")
p("The early chat widget could only offer a few fixed buttons — “Get a quote,” “Learn about "
  "solar,” and so on. The instruction for the finale was simple and human:")
callout("user", "“The chatbot only gives a few options. I want it to be open-ended, dynamic, "
        "smart and AI-based, and answer from our database. We’ve prepared a web chatbot too — "
        "please incorporate it.”")
p("That one sentence set the goal: a free-typing conversation, powered by AI, answering from the "
  "company’s own knowledge — not a menu of canned replies.")

step("Upgrade the chat to an open-ended AI conversation",
     "The widget was rebuilt around a normal message box. A visitor can now type anything and "
     "get an intelligent answer drawn from the company’s knowledge, delivered through the n8n "
     "automation behind the scenes — while the quick “Request a quote” form was kept for "
     "convenience.")
callout("note", "The key safety idea: the browser only ever talks to the company’s own website, "
        "which then quietly relays the question to the AI service. The secret password to the AI "
        "service stays on the server and is never exposed to visitors.")

p("Then came the second instruction:")
callout("user", "“I want the chatbot to process images too.”")

step("Add image understanding to the chatbot",
     "An attach button was added. A visitor can now send a photo — of an electrical panel, an "
     "energy bill, or a rooftop — and the AI can look at it and respond. The picture is "
     "automatically shrunk inside the browser before sending, so even a large phone photo "
     "uploads quickly and reliably.")
figure("fig_chatbot_flow.png", "Figure 6.1 — How a question or photo travels to the AI and back.")
callout("why", "For an engineering firm this is powerful: a prospect can simply photograph their "
        "existing setup and ask “can you upgrade this?” instead of struggling to describe it in "
        "words. The barrier to starting a conversation drops to almost nothing.")
callout("note", "Before sending, the photo is resized and lightly compressed in the visitor’s "
        "browser. An 8-megabyte phone snapshot becomes a few hundred kilobytes — fast even on a "
        "weak mobile connection, with no noticeable loss for understanding the image.")

step("Unify everything under one robust setting",
     "A final tidy-up made sure that however the control panel is configured, the one polished, "
     "brand-styled, image-capable assistant is the one that appears — with no risk of an old "
     "version resurfacing after a future update.")
callout("tip", "When two versions of a feature exist, choose one on purpose and remove the other. "
        "Ambiguity is how the wrong version sneaks back into production months later.")

p("With that, the project was complete: a fast, beautiful, machine-readable website, fully run "
  "from a control panel, live on its own secure server, with an AI assistant that can read both "
  "words and pictures and answer from the company’s own knowledge.")
callout("user", "“Great!!! It works nicely.”")
pagebreak()

# ── PART VII — Publishing pack ───────────────────────────────────────────────
h1("Part VII — The Publishing Pack")
p("This part contains everything needed to turn this manuscript into a published book on "
  "Amazon. Work through the appendices in order.")

h2("Appendix A · Publishing on Amazon, Step by Step")
figure("fig_kdp_steps.png", "Figure A.1 — The path from this file to a live Amazon listing.")
numbered("Finish the manuscript. Add the author photo (save it as ebook/figures/author-photo.jpg "
         "and rebuild, or paste it onto the About-the-Author page), drop in the real screenshots "
         "listed in Appendix E, then update the Table of Contents (right-click it → Update Field → "
         "Update entire table). The author name, bio, and ISBN layout are already in place.")
numbered("Create a free KDP account at kdp.amazon.com using your Amazon login and tax details.")
numbered("Click “Create” and choose a Kindle eBook, a Paperback, and a Hardcover (you can do all "
         "three from one title).")
numbered("Enter the book details: title, subtitle, author, description, keywords, and categories "
         "(all provided in Appendix C).")
numbered("Choose your ISBN for each print format (see Appendix B), or accept a free KDP ISBN.")
numbered("Upload the manuscript file and the cover. For print, export this document to PDF first "
         "(File → Save As → PDF).")
numbered("Use KDP’s previewer to check every page, set your price, and click Publish. Review "
         "usually takes up to 72 hours.")
callout("tip", "Publish the Kindle eBook first to go live fastest, then add paperback and "
        "hardcover. The same title page can host all three formats for buyers to choose.")

h2("Appendix B · ISBN Rules and Standards")
p("An ISBN (International Standard Book Number) is the unique 13-digit fingerprint of a book. "
  "The rules below are the ones that matter for this title.")
h3("Every format needs its own ISBN")
bullet("Kindle eBook: does NOT use an ISBN. Amazon assigns a free ASIN automatically. You may add "
       "your own ISBN if you wish, but it is optional.")
bullet("Paperback: needs its own ISBN, separate from every other format.")
bullet("Hardcover: needs a third, separate ISBN.")
callout("note", "One book, three formats = up to three different ISBNs. An ISBN identifies a "
        "specific format of a specific edition — a new edition or a new format always takes a new "
        "number. The same ISBN is never reused for a different book or format.")
h3("Where to get ISBNs")
bullet("Free from Amazon KDP: fast and costs nothing, but lists Amazon as the technical publisher "
       "and only works for books sold through KDP.")
bullet("Bought from the official agency in your country (for example, Bowker in the USA, or your "
       "national ISBN agency): you appear as the publisher of record and can sell anywhere. In "
       "Bangladesh, ISBNs are issued by the national agency under the Bangladesh National Library.")
callout("why", "If you want “Sustech Technology Ltd” to be the named publisher on every store — "
        "not just Amazon — buy your own ISBNs from your national agency. If you only sell on "
        "Amazon and want it free and simple, the KDP ISBN is fine.")
h3("Printing the ISBN correctly")
bullet("Print the 13-digit ISBN on the copyright page (already laid out in this manuscript).")
bullet("On the print back cover, the ISBN appears inside the barcode, bottom-right. KDP can add "
       "this barcode automatically.")
bullet("Format it with hyphens in the standard groups, e.g. 978-984-XXXXX-X-X.")
callout("note", "Newer ISBNs often begin with 979-8 (common for KDP) or 978. Bangladesh’s country "
        "group inside an ISBN is 984 — so a Bangladeshi-agency ISBN typically reads 978-984-…")

h2("Appendix C · Book Metadata Sheet (copy-paste into KDP)")
h3("Title and subtitle")
config("Title:    Build a World-Class Website with AI\n"
       "Subtitle: The Non-Coder’s Guide — From Idea to AI Chatbot, Step by Step")
h3("Book description (back-cover / Amazon listing)")
p("Can you build a premium, modern company website — one that customers admire and AI "
  "assistants recommend — without writing a single line of code? This book proves you can. "
  "Follow the complete, true story of building the Sustech Technology website from an empty "
  "folder to a live site with an AI chatbot that reads photos. In plain language, with no "
  "jargon left unexplained, you’ll see every decision a real project makes — and learn to "
  "direct AI tools to do the heavy lifting while you stay in charge. If you’ve ever been told "
  "“you need a developer for that,” this book is your answer.", size=10.5)
h3("Seven keywords")
bullet("AI website building for beginners")
bullet("no-code website guide")
bullet("build a website with AI")
bullet("AI for non-programmers")
bullet("small business website")
bullet("AI chatbot for website")
bullet("CMS website step by step")
h3("Two categories")
bullet("Computers & Technology › Web Development & Design")
bullet("Business & Money › Marketing & Sales › Web Marketing")

h2("Appendix D · Cover Design Brief")
p("The cover image included in this manuscript is a working mock-up in the brand palette. For "
  "the final cover, give a designer (or an AI image tool) this brief:")
bullet("Trim size: 6 × 9 inches. Front-cover image ratio for Kindle: 1.6:1 (e.g. 1600 × 2560 px).")
bullet("Palette: True Blue #0073CF, Lime Green #32CD32, Golden Poppy #FCC200, deep ink #0B1B2B.")
bullet("Title dominant at the top; subtitle beneath; author name and “Published by Sustech "
       "Technology Ltd” at the bottom.")
bullet("Mood: confident, premium, engineering-meets-AI. Clean, lots of breathing room, no clutter.")
bullet("For paperback and hardcover, the designer also needs a full wrap (back cover + spine); "
       "KDP provides exact spine width once the page count is final.")

h2("Appendix E · Screenshot Shot-List")
p("Drop these real screenshots into the manuscript where the matching boxes appear, to make the "
  "book vivid. Capture them from the live site and the control panel.")
screenshot_placeholder("E.1", "The finished home page (top section with hero image).", "beta.sustechltd.com")
screenshot_placeholder("E.2", "The control panel editing a page’s blocks.", "beta.sustechltd.com/admin")
screenshot_placeholder("E.3", "A service detail page with its explainer video.", "a /services page")
screenshot_placeholder("E.4", "The AI chatbot open, mid-conversation, with a photo attached.", "any page, chat widget")
callout("tip", "Capture screenshots at a consistent width and on a clean screen (no personal "
        "browser tabs visible). Consistency makes a book look professionally produced.")

h2("Appendix F · Plain-English Glossary")
def gloss(term, mean):
    par = doc.add_paragraph()
    r = par.add_run(term + " — "); r.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(10.5)
    r2 = par.add_run(mean); r2.font.size = Pt(10.5)

gloss("CMS", "Content Management System: the private admin panel where you edit your site’s content without code.")
gloss("Deploy", "To publish the website onto a live server so the public can reach it.")
gloss("GEO / AEO", "Optimising a site so AI engines and answer-boxes can read and recommend it.")
gloss("Hero", "The large headline image or video at the top of a page.")
gloss("ISBN", "The unique 13-digit identifier of a specific book format.")
gloss("n8n", "A visual automation tool; here it powers the chatbot’s connection to the AI.")
gloss("Schema", "A structured fact sheet attached to a page so machines understand it.")
gloss("Server-rendered", "The page’s words are prepared on the server before reaching the browser — fast and machine-readable.")
gloss("VPS", "Virtual Private Server: a rented online computer that runs your website 24/7.")

h2("Appendix G · The Complete Step Index")
p(f"The build was completed in {STEP['n']} counted steps across Parts II–VI, each one a real, "
  "final result with reworks and dead-ends omitted. Use the numbered steps as a checklist if you "
  "set out to build your own site along the same path.")

h2("About the Author")
# Author photo (auto-embeds if ebook/figures/author-photo.jpg is present)
_author_photo = None
for _cand in ("author-photo.jpg", "author-photo.jpeg", "author-photo.png"):
    _pp = os.path.join(FIG, _cand)
    if os.path.exists(_pp):
        _author_photo = _pp
        break
if _author_photo:
    _ap = doc.add_paragraph(); _ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ap.add_run().add_picture(_author_photo, width=Inches(2.2))
else:
    screenshot_placeholder("Author", "Author portrait — save your photo as "
                           "ebook/figures/author-photo.jpg and rebuild.", "About the Author page")
p("Md. Sohel Sikder", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=INK, bold=True)
p("Managing Director & Founder, Sustech Technology Ltd",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=BLUE, italic=True)
p("Md. Sohel Sikder is a CUET graduate, a SREDA-certified Energy Auditor, and a "
  "Government-certified Fire Safety Manager. He founded Sustech Technology Ltd on the conviction "
  "that Bangladesh’s industrial sector deserved engineering partners who combine global standards "
  "with deep local expertise — leading the company’s strategy, key accounts, EPC oversight, and "
  "major technical proposals.")
p("This book grew out of a real project: rebuilding the Sustech website by directing modern AI "
  "tools rather than writing code by hand — proof that a determined non-programmer can ship a "
  "premium, AI-ready website. He shares it so other founders and professionals can do the same.")
p("Published by Sustech Technology Ltd — an engineering firm delivering solar, electrical, "
  "lightning-protection, and smart-systems projects for industry in Bangladesh.")

# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print("Saved:", OUT)
print("Total counted build steps:", STEP["n"])
